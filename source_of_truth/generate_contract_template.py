"""Generate contract-template.docx from the annotated HTML source.

This script extracts the German contract text (including placeholders) from
`contract-template-annotated.html` and writes a simplified DOCX version.
It omits styling/annotations while preserving headings and line breaks so the
backend can populate placeholders.
"""
from pathlib import Path
from typing import Iterable

from docx import Document
from lxml import html

ROOT = Path(__file__).resolve().parent
ANNOTATED_HTML = ROOT / "contract-template-annotated.html"
OUTPUT_DOCX = ROOT / "contract-template.docx"


def _clean_text(text: str) -> str:
    """Normalize whitespace while preserving intentional blank lines."""
    lines = [line.strip() for line in text.splitlines()]
    cleaned: list[str] = []
    for line in lines:
        if line:
            cleaned.append(line)
        elif cleaned and cleaned[-1] != "":
            cleaned.append("")
    while cleaned and cleaned[-1] == "":
        cleaned.pop()
    return "\n".join(cleaned)


def _iter_german_blocks(section) -> Iterable[str]:
    for block in section.xpath(".//div[contains(@class, 'german')]"):
        raw_text = block.text_content()
        cleaned = _clean_text(raw_text)
        if cleaned:
            yield cleaned


def build_docx() -> None:
    tree = html.parse(str(ANNOTATED_HTML))
    doc = Document()

    doc.add_heading("Wohnraummietvertrag (Template)", level=1)
    doc.add_paragraph(
        "Automated export generated from contract-template-annotated.html. "
        "Replace placeholders like [LANDLORD_NAME] programmatically."
    )

    for section in tree.xpath("//div[@class='section']"):
        titles = section.xpath(".//div[contains(@class, 'section-title')]/text()")
        if titles:
            doc.add_heading(titles[0].strip(), level=2)

        for block_text in _iter_german_blocks(section):
            for paragraph in block_text.split("\n\n"):
                doc.add_paragraph(paragraph)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(f"Wrote {OUTPUT_DOCX.relative_to(Path.cwd())}")
