import io
import json
import logging
import os
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, Iterable, Optional, Tuple
from uuid import uuid4

import azure.functions as func
from azure.core.exceptions import AzureError
from azure.storage.blob import BlobClient, BlobServiceClient
from docx import Document
from pydantic import BaseModel, ValidationError

from domain.contract_context import build_contract_context, ContractContextError


PLACEHOLDER_PATTERN = re.compile(r"\[([A-Z0-9_]+)\]")


# ============================================================
# Errors
# ============================================================


class TemplateProcessingError(RuntimeError):
    pass


# ============================================================
# Storage Resolution (CRITICAL FIX)
# ============================================================


def _clean(value: Optional[str]) -> Optional[str]:
    if value and value.strip():
        return value.strip()
    return None


def _resolve_connection(*env_vars: str, default: str = "UseDevelopmentStorage=true") -> str:
    for var in env_vars:
        candidate = _clean(os.getenv(var))
        if candidate:
            return candidate
    return default


@dataclass
class StorageSettings:
    template_path: str
    template_container: str
    template_connection: str
    contracts_container: str
    contracts_connection: str

    @classmethod
    def from_env(cls, template_path_override: Optional[str]) -> "StorageSettings":
        template_path = _clean(template_path_override) or _clean(
            os.getenv("TemplateBlobPath")
        )

        if not template_path:
            raise TemplateProcessingError(
                "Template path not provided. Supply 'templatePath' in the request or set the TemplateBlobPath environment variable."
            )

        return cls(
            template_path=template_path,
            template_container=_clean(os.getenv("TemplatesContainer")) or "templates",
            template_connection=_resolve_connection(
                "TemplateBlobConnection", "ContractsBlobConnection", "AzureWebJobsStorage"
            ),
            contracts_container=_clean(os.getenv("ContractsContainer")) or "contracts",
            contracts_connection=_resolve_connection(
                "ContractsBlobConnection", "AzureWebJobsStorage"
            ),
        )

    @staticmethod
    def template_hint() -> str:
        return (
            "Check TemplateBlobConnection, ContractsBlobConnection, AzureWebJobsStorage, and TemplatesContainer environment values."
        )

    @staticmethod
    def contract_hint() -> str:
        return "Verify ContractsBlobConnection, AzureWebJobsStorage, and ContractsContainer environment values."


# ============================================================
# Request Model
# ============================================================


class GenerateContractRequest(BaseModel):
    maskA: Dict[str, Any]
    maskB: Dict[str, Any]
    templatePath: Optional[str] = None

    class Config:
        extra = "forbid"


# ============================================================
# Helpers
# ============================================================


def _error_response(message: str, status_code: int) -> func.HttpResponse:
    return func.HttpResponse(
        json.dumps({"message": message}, ensure_ascii=False),
        status_code=status_code,
        mimetype="application/json",
    )


def _load_template(
    template_path: str, connection_string: str, container: str
) -> Tuple[str, bytes]:
    # Local file
    if os.path.exists(template_path):
        with open(template_path, "rb") as f:
            return os.path.splitext(template_path)[1].lower(), f.read()

    if not connection_string:
        raise TemplateProcessingError(
            f"Template storage connection not configured for '{template_path}'. {StorageSettings.template_hint()}"
        )

    blob_name = template_path.lstrip("/")
    container_name = container

    if "/" in blob_name:
        maybe_container, remainder = blob_name.split("/", 1)
        if remainder:
            container_name = maybe_container
            blob_name = remainder

    try:
        service = BlobServiceClient.from_connection_string(connection_string)
        container_client = service.get_container_client(container_name)

        if not container_client.exists():
            raise TemplateProcessingError(
                f"Template container '{container_name}' not found while attempting to read '{template_path}'. {StorageSettings.template_hint()}"
            )

        blob_client = container_client.get_blob_client(blob_name)

        if not blob_client.exists():
            raise TemplateProcessingError(
                f"Template blob '{blob_name}' not found in container '{container_name}'. {StorageSettings.template_hint()}"
            )

        data = blob_client.download_blob().readall()

    except TemplateProcessingError:
        raise
    except (AzureError, ValueError) as exc:
        raise TemplateProcessingError(
            f"Failed to load template '{template_path}' from blob storage. {StorageSettings.template_hint()}"
        ) from exc

    return os.path.splitext(template_path)[1].lower(), data


def _apply_placeholders_to_runs(
    runs: Iterable[Any], placeholders: Dict[str, str], used: set[str]
) -> None:
    for run in runs:
        for key, value in placeholders.items():
            marker = f"[{key}]"
            if marker in run.text:
                run.text = run.text.replace(marker, value)
                used.add(key)


def _extract_placeholders(text: str) -> set[str]:
    return set(PLACEHOLDER_PATTERN.findall(text or ""))


def _collect_template_placeholders(document: Document) -> set[str]:
    found: set[str] = set()

    for paragraph in document.paragraphs:
        found.update(_extract_placeholders(paragraph.text))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                found.update(_extract_placeholders(cell.text))

    return found


def _render_docx_template(template_bytes: bytes, context: Dict[str, str]) -> bytes:
    document = Document(io.BytesIO(template_bytes))
    used: set[str] = set()

    template_placeholders = _collect_template_placeholders(document)
    context_keys = set(context.keys())

    missing = template_placeholders - context_keys
    if missing:
        raise TemplateProcessingError(
            "Missing placeholders in context: " + ", ".join(sorted(missing))
        )

    unexpected = context_keys - template_placeholders
    if unexpected:
        raise TemplateProcessingError(
            "Unexpected placeholders not in template: " + ", ".join(sorted(unexpected))
        )

    for p in document.paragraphs:
        _apply_placeholders_to_runs(p.runs, context, used)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply_placeholders_to_runs(p.runs, context, used)

    missing = set(context.keys()) - used
    if missing:
        raise TemplateProcessingError(
            "Unreplaced placeholders in template: " + ", ".join(sorted(missing))
        )

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _render_html_template(html_bytes: bytes, context: Dict[str, str]) -> bytes:
    html = html_bytes.decode("utf-8", errors="ignore")

    for key, value in context.items():
        html = html.replace(f"[{key}]", value)

    unresolved = re.findall(r"\[[A-Z0-9_]+\]", html)
    if unresolved:
        raise TemplateProcessingError(
            "Unreplaced placeholders in template: " + ", ".join(sorted(set(unresolved)))
        )

    text = re.sub(r"<[^>]+>", "", html)

    document = Document()
    for block in filter(None, (b.strip() for b in text.split("\n\n"))):
        document.add_paragraph(block)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _upload_contract(
    contract_bytes: bytes, connection_string: str, container: str
) -> Tuple[str, str]:
    try:
        if not connection_string:
            raise TemplateProcessingError(
                f"Contract storage connection not configured. {StorageSettings.contract_hint()}"
            )

        service = BlobServiceClient.from_connection_string(connection_string)
        container_client = service.get_container_client(container)

        if not container_client.exists():
            container_client.create_container()

        blob_name = f"contract-{datetime.utcnow():%Y%m%d%H%M%S}-{uuid4()}.docx"
        blob_client = container_client.get_blob_client(blob_name)

        blob_client.upload_blob(
            contract_bytes,
            overwrite=True,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return blob_name, blob_client.url

    except TemplateProcessingError:
        raise
    except (AzureError, ValueError) as exc:
        logging.exception("Blob upload failed")
        raise TemplateProcessingError(
            f"Failed to upload generated contract to container '{container}'. {StorageSettings.contract_hint()}"
        ) from exc


# ============================================================
# Azure Function Entry Point
# ============================================================


async def main(req: func.HttpRequest) -> func.HttpResponse:
    try:
        body = req.get_json()
    except ValueError:
        return _error_response("Invalid JSON body.", 400)

    try:
        payload = GenerateContractRequest(**body)
    except ValidationError as exc:
        return func.HttpResponse(exc.json(), status_code=400, mimetype="application/json")

    # Build contract context
    try:
        context = build_contract_context(payload.maskA, payload.maskB)
    except ContractContextError as exc:
        return _error_response(str(exc), 400)

    try:
        storage_settings = StorageSettings.from_env(payload.templatePath)
    except TemplateProcessingError as exc:
        return _error_response(str(exc), 400)

    # Load template
    try:
        extension, template_bytes = _load_template(
            storage_settings.template_path,
            storage_settings.template_connection,
            storage_settings.template_container,
        )
    except TemplateProcessingError as exc:
        return _error_response(str(exc), 400)

    # Render contract
    try:
        if extension == ".docx":
            contract_bytes = _render_docx_template(template_bytes, context)
        else:
            contract_bytes = _render_html_template(template_bytes, context)
    except TemplateProcessingError as exc:
        return _error_response(str(exc), 400)

    # Upload
    try:
        _, download_url = _upload_contract(
            contract_bytes,
            storage_settings.contracts_connection,
            storage_settings.contracts_container,
        )
    except TemplateProcessingError as exc:
        return _error_response(str(exc), 500)

    return func.HttpResponse(
        json.dumps({"downloadUrl": download_url}, ensure_ascii=False),
        status_code=200,
        mimetype="application/json",
    )

