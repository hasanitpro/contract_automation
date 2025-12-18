import io
import os
import socket
import subprocess
import time
from pathlib import Path
import sys

import pytest
from azure.core.exceptions import ResourceExistsError
from azure.storage.blob import BlobServiceClient
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import generate_contract as gc  # noqa: E402


DEVSTORE_KEY = "Eby8vdM02xNOcqFlqUwJPLlmEtlCDXJ1OUzFT50uSRZ6IFsuFq2UVErCz4I6jAz=="


def _find_free_port() -> int:
    with socket.socket() as sock:
        sock.bind(("127.0.0.1", 0))
        return sock.getsockname()[1]


@pytest.fixture(scope="session")
def azurite(tmp_path_factory):
    data_dir = tmp_path_factory.mktemp("azurite-data")
    port = _find_free_port()
    queue_port = _find_free_port()
    table_port = _find_free_port()
    process: subprocess.Popen | None = None
    os.environ.setdefault("AZURE_STORAGE_API_VERSION", "2021-08-06")
    try:
        subprocess.run(
            ["npx", "--yes", "azurite", "--version"],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    except FileNotFoundError:
        pytest.skip("Azurite (npx) is required for these tests")
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(f"Failed to prepare Azurite: {exc}") from exc
    cmd = [
        "npx",
        "--yes",
        "azurite",
        "--silent",
        "--location",
        str(data_dir),
        "--blobHost",
        "127.0.0.1",
        "--blobPort",
        str(port),
        "--queueHost",
        "127.0.0.1",
        "--queuePort",
        str(queue_port),
        "--tableHost",
        "127.0.0.1",
        "--tablePort",
        str(table_port),
    ]

    try:
        process = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            env={**os.environ, "AZURITE_ACCOUNTS": f"devstoreaccount1:{DEVSTORE_KEY}"},
        )
    except FileNotFoundError:
        pytest.skip("Azurite (npx) is required for these tests")

    try:
        deadline = time.time() + 30
        while time.time() < deadline:
            with socket.socket() as sock:
                if sock.connect_ex(("127.0.0.1", port)) == 0:
                    break
            time.sleep(0.25)
        else:
            if process and process.poll() is None:
                process.terminate()
            stdout, stderr = process.communicate(timeout=5)
            raise RuntimeError(
                "Azurite did not start in time: "
                f"stdout={stdout.decode(errors='ignore')} stderr={stderr.decode(errors='ignore')}"
            )

        connection_string = (
            "DefaultEndpointsProtocol=http;"
            "AccountName=devstoreaccount1;"
            f"AccountKey={DEVSTORE_KEY};"
            f"BlobEndpoint=http://127.0.0.1:{port}/devstoreaccount1;"
            f"QueueEndpoint=http://127.0.0.1:{queue_port}/devstoreaccount1;"
            f"TableEndpoint=http://127.0.0.1:{table_port}/devstoreaccount1;"
        )

        yield connection_string
    finally:
        if process:
            process.terminate()
            try:
                process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                process.kill()


@pytest.fixture
def sample_docx_bytes():
    doc = Document()
    doc.add_paragraph("Hello [NAME]")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_load_template_from_azurite(azurite, sample_docx_bytes):
    service = BlobServiceClient.from_connection_string(azurite)
    container = service.get_container_client("templates")
    try:
        container.create_container()
    except ResourceExistsError:
        pass
    container.upload_blob("contract.docx", sample_docx_bytes, overwrite=True)

    extension, data = gc._load_template("contract.docx", azurite, "templates")

    assert extension == ".docx"
    assert data[:4] == sample_docx_bytes[:4]


def test_load_template_missing_blob_reports_hint(azurite):
    service = BlobServiceClient.from_connection_string(azurite)
    container = service.get_container_client("templates")
    try:
        container.create_container()
    except ResourceExistsError:
        pass

    with pytest.raises(gc.TemplateProcessingError) as excinfo:
        gc._load_template("missing.docx", azurite, "templates")

    message = str(excinfo.value)
    assert "Template blob 'missing.docx'" in message
    assert "TemplatesContainer" in message


def test_render_docx_template(sample_docx_bytes):
    rendered = gc._render_docx_template(sample_docx_bytes, {"NAME": "Alice"})

    doc = Document(io.BytesIO(rendered))
    assert any("Alice" in paragraph.text for paragraph in doc.paragraphs)


def test_render_docx_template_reports_unused_placeholders(sample_docx_bytes):
    with pytest.raises(gc.TemplateProcessingError) as excinfo:
        gc._render_docx_template(sample_docx_bytes, {"NAME": "Alice", "EXTRA": "value"})

    assert "Unreplaced placeholders" in str(excinfo.value)


def test_upload_contract_to_azurite(azurite, sample_docx_bytes):
    blob_name, url = gc._upload_contract(sample_docx_bytes, azurite, "contracts")

    service = BlobServiceClient.from_connection_string(azurite)
    blob_client = service.get_container_client("contracts").get_blob_client(blob_name)

    assert blob_client.exists()
    assert blob_client.download_blob().readall()
    assert url.endswith(blob_name)


def test_upload_contract_bad_connection_reports_hint(monkeypatch, sample_docx_bytes):
    def raise_error(_conn):
        raise gc.AzureError("boom")

    monkeypatch.setattr(gc.BlobServiceClient, "from_connection_string", raise_error)

    with pytest.raises(gc.TemplateProcessingError) as excinfo:
        gc._upload_contract(sample_docx_bytes, "AccountEndpoint=https://example.blob.core.windows.net/", "contracts")

    message = str(excinfo.value)
    assert "Failed to upload generated contract" in message
    assert "ContractsContainer" in message
